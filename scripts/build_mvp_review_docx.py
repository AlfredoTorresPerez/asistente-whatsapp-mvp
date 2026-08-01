from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "outputs" / "revision_mvp_word"
OUTPUT_PATH = OUTPUT_DIR / "Revision_Comite_MVP_Asistente_WhatsApp.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
RISK_RED = RGBColor(155, 28, 28)
GREEN = RGBColor(22, 101, 52)
GOLD = RGBColor(122, 90, 0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_title(doc: Document, text: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run, size=12.5, color=MUTED)


def add_metadata(doc: Document) -> None:
    rows = [
        ("Proyecto", "Asistente empresarial para WhatsApp - Centro Estetico Bella"),
        ("Tipo de revision", "Comite tecnico y funcional previo a demostracion cliente"),
        ("Fecha de evaluacion", "04/06/2026"),
        ("Veredicto conjunto", "El MVP solo debe mostrarse en una demostracion controlada."),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.7, 4.65])
    for i, (label, value) in enumerate(rows):
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, color=INK, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_run_font(r, size=10.8, color=INK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_run_font(r, size=10.8, color=INK)


def add_callout(doc: Document, label: str, text: str, fill: str = "F8FAFC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p.add_run("\n")
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc: Document) -> None:
    headers = ["Area", "Estado", "Que esta bien", "Que esta mal", "Que falta", "Riesgo", "Prioridad"]
    rows = [
        ["Funcional", "Riesgoso", "Login, dashboard, conversaciones, administracion, IA y canal de WhatsApp existen.",
         "No todo esta validado end-to-end hoy.", "Guion demo y pruebas manuales cerradas.", "Medio", "Alta"],
        ["Tecnico", "Riesgoso", "Docker levanta healthy; backend inicia; DB migra a v14.",
         "Tests backend fallan; frontend build/test/lint fallan.", "tsconfig.json, config Vitest, limpieza lint.", "Alto", "Alta"],
        ["Conversacional", "Riesgoso", "Hay agentes IA, contexto, reglas DB, handoff y auditoria.",
         "Test de coherencia IA falla por normalizacion de acentos.", "Mas pruebas de agenda/precios/servicios.", "Medio", "Alta"],
        ["Demostracion", "Correcto con restricciones", "Stack responde en 8080 y 5173; logs filtran mensajes inutiles y propios.",
         "No conviene improvisar con IA automatica o QR no preparado.", "Checklist antes de mostrar.", "Medio", "Alta"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, [0.85, 1.05, 1.35, 1.35, 1.25, 0.62, 0.62])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=8.4, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=8.2, color=INK)
            if i in (0, 1, 5, 6):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.3, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Revision MVP WhatsApp Empresarial")
    set_run_font(r, size=9, color=MUTED)

    add_title(
        doc,
        "Revision Comite MVP Asistente WhatsApp Empresarial",
        "Evaluacion tecnica y funcional previa a demostracion cliente",
    )
    add_metadata(doc)

    add_heading(doc, "1. Resumen ejecutivo")
    add_callout(
        doc,
        "Decision final",
        "El MVP solo debe mostrarse en una demostracion controlada. El producto comunica valor y el stack local levanta, pero no esta listo para produccion ni para una demostracion libre sin guion.",
        fill="FFF7ED",
    )
    add_para(doc, "Estado general: el MVP funciona como demo local controlada, pero no esta listo para produccion.")
    add_para(doc, "Nivel de preparacion para cliente: medio. Se puede mostrar si se controla el recorrido y se evitan promesas de automatizacion completa.")
    add_para(doc, "Principales fortalezas: stack Docker healthy, backend Spring Boot activo, PostgreSQL con migraciones aplicadas hasta v14, frontend disponible en localhost:5173, modulo IA del Negocio, administracion, conversaciones, catalogo, agenda y canal de WhatsApp (simulado o Cloud API).")
    add_para(doc, "Principales debilidades: backend tests fallan, frontend build falla por ausencia de tsconfig.json, frontend tests fallan por configuracion Vitest/jsdom, lint falla con errores React Hooks, el canal WhatsApp depende de simulador o Cloud API de Meta.")

    add_heading(doc, "2. Matriz de evaluacion")
    add_matrix(doc)

    add_heading(doc, "3. Evidencia revisada")
    add_simple_table(
        doc,
        ["Evidencia", "Resultado", "Lectura"],
        [
            ["Docker Compose local", "Servicios backend, frontend y postgres healthy.", "Apto para demo local tras levantar el stack."],
            ["HTTP backend", "GET /api/v1/health respondio UP.", "Backend disponible."],
            ["HTTP frontend", "localhost:5173 respondio 200.", "Interfaz disponible."],
            ["Flyway", "14 migraciones validadas, schema en version 14.", "Base de datos consistente en el entorno actual."],
            ["Logs del canal WhatsApp", "Ignora notification_template y mensajes salientes propios.", "Reduce ruido y riesgo de duplicidad."],
            ["Backend tests", "14 tests ejecutados; 1 falla en coherencia IA.", "No esta limpio para release."],
            ["Frontend build", "Falla por ausencia de tsconfig.json.", "Build productivo no verificable localmente."],
            ["Frontend tests", "Fallan por document/describe no definidos.", "Vitest requiere configuracion jsdom/globals."],
            ["Frontend lint", "5 errores y 14 warnings.", "Deuda tecnica antes de produccion."],
        ],
        [1.55, 2.25, 2.55],
    )

    add_heading(doc, "4. Que esta bien")
    add_bullets(doc, [
        "El stack Docker local puede levantarse y quedar healthy.",
        "Backend Spring Boot inicia correctamente y expone health check.",
        "PostgreSQL valida migraciones y conserva el esquema en version v14.",
        "Frontend responde y permite preparar una demostracion visual.",
        "Existe modulo IA del Negocio y rutas principales de administracion.",
        "El canal WhatsApp tiene filtros para mensajes inutiles y mensajes propios.",
        "La auto-respuesta IA esta desactivada por defecto, lo cual reduce riesgo en demo.",
        "El dominio del centro estetico esta modelado con servicios, catalogo, agenda y conversaciones.",
    ])

    add_heading(doc, "5. Que esta mal")
    add_bullets(doc, [
        "El backend no pasa la suite completa por una prueba de coherencia IA.",
        "El frontend no puede ejecutar build productivo local porque falta tsconfig.json.",
        "Los tests frontend no cargan entorno de navegador ni globals de Vitest.",
        "El lint frontend falla por patrones React Hooks en catalogo, conversaciones y reglas.",
        "El contenedor frontend usa Vite dev server; no representa un despliegue productivo real.",
        "Los logs de Spring muestran password generado de desarrollo; debe revisarse antes de produccion.",
        "El canal WhatsApp via Cloud API de Meta es el camino productivo; el simulador queda solo para desarrollo.",
    ])

    add_heading(doc, "6. Que falta")
    add_bullets(doc, [
        "Agregar tsconfig.json y confirmar pnpm build exitoso.",
        "Configurar Vitest con jsdom, globals y setupFiles.",
        "Corregir errores de lint React Hooks.",
        "Resolver la prueba de coherencia IA o normalizar la comparacion de acentos.",
        "Agregar pruebas end-to-end del flujo WhatsApp: recibir, clasificar, sugerir y responder manualmente.",
        "Validar QR real o preparar simulacion estable antes de la reunion.",
        "Documentar checklist de demostracion con mensajes exactos y rutas a mostrar.",
        "Separar claramente modo demo, piloto y produccion.",
    ])

    add_heading(doc, "7. Riesgos antes de mostrar al cliente")
    add_simple_table(
        doc,
        ["Nivel", "Riesgos", "Mitigacion"],
        [
            ["Critico", "Mostrarlo como produccion o activar auto-respuesta real sin validacion.", "Presentarlo solo como demo controlada; mantener auto-reply apagado."],
            ["Critico", "Prometer estabilidad del canal simulado.", "Explicar que para produccion se recomienda WhatsApp Cloud API."],
            ["Medio", "Cliente tecnico pregunta por QA y aparecen tests fallidos.", "Tener respuesta honesta: la demo valida valor, la etapa siguiente limpia CI y pruebas."],
            ["Medio", "QR o sesion WhatsApp no esta lista durante la reunion.", "Preparar simulacion o verificar sesion antes de iniciar."],
            ["Menor", "Logs de mensajes ignorados parecen errores.", "Explicar que son filtros esperados del adaptador."],
        ],
        [0.9, 3.0, 2.45],
    )

    add_heading(doc, "8. Correcciones obligatorias antes de mostrar")
    add_numbered(doc, [
        "Verificar docker compose ps y health checks antes de la reunion.",
        "Probar login demo y navegacion privada.",
        "Probar Dashboard, Conversaciones, IA del Negocio, Administracion y el canal de WhatsApp.",
        "Preparar QR o simulacion; no improvisar conexion real.",
        "Mantener APP_AI_AGENTS_AUTO_REPLY_ENABLED=false.",
        "Corregir o documentar internamente el test de coherencia IA fallido.",
    ])

    add_heading(doc, "9. Correcciones recomendadas despues de mostrar")
    add_bullets(doc, [
        "Arreglar build frontend productivo.",
        "Arreglar configuracion Vitest.",
        "Limpiar lint y agregar CI obligatorio.",
        "Agregar pruebas de controladores administrativos y flujo multiagente.",
        "Migrar canal productivo a WhatsApp Cloud API.",
        "Mejorar auditoria visible de decisiones IA.",
        "Agregar monitoreo, backups, gestion de secretos y hardening de seguridad.",
    ])

    add_heading(doc, "10. Flujo recomendado para demostracion")
    add_simple_table(
        doc,
        ["Bloque", "Mostrar", "Evitar"],
        [
            ["Inicio", "Login demo, Dashboard y KPIs.", "Explicar como si ya fuera produccion."],
            ["Operacion", "Conversaciones, historial, envio manual y estado visual.", "Mensajes reales no preparados."],
            ["IA", "IA del Negocio, respuestas sugeridas, auditoria y modo supervisado.", "Auto-respuesta real sin aprobacion humana."],
            ["Centro estetico", "Catalogo, servicios, precios y agenda como contexto comercial.", "Confirmar disponibilidad sin validacion."],
            ["Administracion", "Usuarios, seguridad y canal de WhatsApp.", "Prometer SLA de WhatsApp Web experimental."],
        ],
        [1.0, 2.75, 2.6],
    )
    add_para(doc, "Mensajes sugeridos para prueba: 'Hola, quiero saber precios de depilacion.', 'Quiero agendar depilacion bozo manana a las 14:00.', 'Quiero hablar con una persona.', 'Que servicios tienen?'.")

    add_heading(doc, "11. Veredictos de los agentes")
    add_callout(
        doc,
        "Orquestador WhatsApp Empresarial",
        "Aprobado con restricciones. El MVP comunica valor comercial y permite mostrar atencion, gestion, IA, conversaciones y operacion de centro estetico, pero debe mostrarse solo como demo controlada.",
        fill="EFF6FF",
    )
    add_callout(
        doc,
        "Asistente WhatsApp Empresarial Tecnico",
        "Tecnicamente listo con riesgos. El stack corre y los servicios estan healthy, pero las pruebas y el build local no estan limpios. Para produccion es insuficiente; para demo guiada es aceptable.",
        fill="F8FAFC",
    )
    add_callout(
        doc,
        "Veredicto final conjunto",
        "El MVP solo debe mostrarse en una demostracion controlada.",
        fill="FEF3C7",
    )

    add_heading(doc, "12. Respuesta sugerida ante preguntas del cliente")
    add_simple_table(
        doc,
        ["Pregunta posible", "Respuesta recomendada"],
        [
            ["Esto ya responde automaticamente en produccion?", "En esta demo esta en modo supervisado. La auto-respuesta se activa despues de validar calidad, agenda y reglas de negocio."],
            ["WhatsApp Web es el canal definitivo?", "El simulador cubre demo y desarrollo; para produccion se usa WhatsApp Cloud API de Meta."],
            ["Puede confirmar horas automaticamente?", "Puede guiar la agenda, pero la confirmacion final debe validar disponibilidad real antes de comprometer horario."],
            ["La IA inventa precios o servicios?", "El objetivo es responder desde catalogo y reglas del negocio; antes de produccion se amplian pruebas para evitar invenciones."],
            ["Que falta para produccion?", "QA automatizado limpio, canal productivo, seguridad, monitoreo, backups, secretos y pruebas end-to-end."],
        ],
        [2.1, 4.25],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "Anexo: checklist rapido previo a la demo")
    add_bullets(doc, [
        "Confirmar que backend, frontend y postgres esten healthy.",
        "Abrir http://localhost:5173 y validar login.",
        "Confirmar que /api/v1/health responde UP.",
        "Revisar que el QR o simulador este preparado.",
        "Mantener auto-respuesta IA desactivada.",
        "Tener mensajes de prueba escritos antes de compartir pantalla.",
        "Evitar prometer produccion sin explicar la fase de hardening.",
    ])

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build()
    print(OUTPUT_PATH)
